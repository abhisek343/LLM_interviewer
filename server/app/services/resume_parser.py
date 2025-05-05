# LLM_interviewer/server/app/services/resume_parser.py

import logging
from pathlib import Path
from typing import Optional

import aiofiles
from pypdf import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

class ResumeParserError(Exception):
    """Custom exception for resume parsing errors."""
    pass

async def parse_resume(file_path: Path) -> str:
    """
    Parses the content of a resume file (PDF or DOCX) and returns the extracted text.

    Args:
        file_path: Path object pointing to the resume file.

    Returns:
        The extracted text content as a string.

    Raises:
        ResumeParserError: If the file type is unsupported or parsing fails.
    """
    file_extension = file_path.suffix.lower()
    logger.info(f"Attempting to parse resume file: {file_path} (type: {file_extension})")

    try:
        if file_extension == ".pdf":
            # pypdf is synchronous, so read file async then parse
            async with aiofiles.open(file_path, "rb") as f:
                pdf_bytes = await f.read()
            import io # Required for BytesIO
            content = ""
            try:
                reader = PdfReader(io.BytesIO(pdf_bytes))
                content_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
                content = "\n".join(content_parts).strip()
                if content:
                    logger.info(f"Successfully parsed PDF with pypdf: {file_path}, extracted {len(content)} characters.")
                    return content
                else: # pypdf succeeded but extracted no content
                    logger.warning(f"pypdf extracted no content from PDF (reader succeeded but no text): {file_path}.")
            except Exception as pypdf_error: # Includes PdfStreamError and other pypdf issues
                logger.warning(f"pypdf failed to process PDF {file_path}: {pypdf_error}. Attempting plain text read as fallback.")

            # Fallback to plain text if pypdf failed or yielded no content
            logger.debug(f"PDF Fallback: pdf_bytes (first 100): {pdf_bytes[:100]}")
            try:
                decoded_text = pdf_bytes.decode('utf-8', errors='ignore')
                logger.debug(f"PDF Fallback: decoded_text (before strip): '{decoded_text}'")
                decoded_text_stripped = decoded_text.strip()
                logger.debug(f"PDF Fallback: decoded_text_stripped: '{decoded_text_stripped}'")
                if decoded_text_stripped:
                    logger.info(f"Successfully read PDF as plain text (fallback): {file_path}, {len(decoded_text_stripped)} chars.")
                    return decoded_text_stripped
                else:
                    logger.warning(f"Plain text fallback also yielded no content for PDF: {file_path} (decoded_text_stripped was empty)")
                    return "" # Return empty string if fallback is also empty
            except Exception as text_e:
                logger.error(f"Error during plain text fallback for PDF {file_path}: {text_e}", exc_info=True)
                return "" # Return empty string on fallback error
        elif file_extension == ".docx":
            # python-docx is synchronous
            # It can take a file path directly.
            # To keep the function async and avoid blocking for large files,
            # one might consider running this in a thread pool executor,
            # but for typical resume sizes, direct call might be acceptable.
            # For simplicity here, direct call. If blocking becomes an issue, refactor.
            try:
                # python-docx Document() can take a path string or a file-like object.
                # Reading with aiofiles first to maintain async pattern, though Document itself is sync.
                async with aiofiles.open(file_path, "rb") as f:
                    doc_bytes = await f.read()
                
                import io # Required for BytesIO
                doc = Document(io.BytesIO(doc_bytes))
                text_parts = [para.text for para in doc.paragraphs if para.text]
                content = "\n".join(text_parts)
                if not content.strip():
                    logger.warning(f"Extracted empty content from DOCX: {file_path}")
                logger.info(f"Successfully parsed DOCX: {file_path}, extracted {len(content)} characters.")
                return content
            except Exception as e: # Catch specific docx errors if known, e.g., PackageNotFoundError
                logger.error(f"Error parsing DOCX file {file_path} with python-docx: {e}", exc_info=True)
                raise ResumeParserError(f"Failed to parse DOCX file: {e}")
        else:
            logger.error(f"Unsupported file type: {file_extension} for file {file_path}")
            raise ResumeParserError(f"Unsupported file type: {file_extension}")
    except ResumeParserError: # Re-raise specific errors
        raise
    except Exception as e:
        logger.error(f"An unexpected error occurred while parsing {file_path}: {e}", exc_info=True)
        raise ResumeParserError(f"Unexpected error parsing file: {e}")
